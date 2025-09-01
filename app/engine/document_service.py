# File: app/engine/document_service.py
# Entity-agnostic document service following Universal Engine patterns
# Backend-heavy, configuration-driven, no assumptions

from flask import session, render_template, request, make_response
from flask_login import current_user
from datetime import datetime
from typing import Dict, Any, Optional, List
import os
import logging

from app.config.entity_configurations import get_entity_config
from app.services.database_service import get_db_session
from app.utils.context_helpers import (
    get_complete_context,
    get_branch_uuid_from_context_or_request
)


logger = logging.getLogger(__name__)

class UniversalDocumentService:
    """
    Entity-agnostic document service for Universal Engine
    Reuses existing data from Universal Views - NO recalculation
    """
    
    def __init__(self):
        """Initialize document service"""
        self.logger = logger
        
    def get_document_data_from_session(self, entity_type: str, entity_id: str) -> Optional[Dict]:
        """
        Retrieve pre-calculated data from session stored by Universal View
        PRINCIPLE: Zero additional queries, reuse Universal Engine data
        """
        session_key = f"doc_data_{entity_type}_{entity_id}"
        
        if session_key in session:
            data = session.get(session_key)
            # Check if data is fresh (within 30 minutes)
            if 'timestamp' in data:
                from datetime import datetime, timedelta
                stored_time = datetime.fromisoformat(data['timestamp'])
                if datetime.now() - stored_time < timedelta(minutes=30):
                    return data.get('data', {})
        
        # If no session data, return None (view will handle fetching)
        return None
    
    def store_document_data_in_session(self, entity_type: str, entity_id: str, data: Dict):
        """
        Store assembled data in session for document generation
        """
        session_key = f"doc_data_{entity_type}_{entity_id}"
        session[session_key] = {
            'data': data,
            'timestamp': datetime.now().isoformat()
        }
        
        # Clean up old session data (keep only last 5 entities)
        self._cleanup_session_data(entity_type)
    
    def _cleanup_session_data(self, entity_type: str):
        """Clean up old session data to prevent session bloat"""
        entity_keys = [k for k in session.keys() if k.startswith(f"doc_data_{entity_type}_")]
        
        if len(entity_keys) > 5:
            # Sort by timestamp and keep only newest 5
            sorted_keys = sorted(
                entity_keys, 
                key=lambda k: session.get(k, {}).get('timestamp', ''),
                reverse=True
            )
            
            for key in sorted_keys[5:]:
                session.pop(key, None)
    
        
    def enhance_data_with_relationships(self, entity_config, data: Dict) -> Dict:
        """
        Enhance data with relationship fields if needed
        Fixed to handle entity_config as dict or object
        """
        # Get include_calculated_fields safely
        include_calculated_fields = []
        if hasattr(entity_config, 'include_calculated_fields'):
            include_calculated_fields = entity_config.include_calculated_fields
        elif isinstance(entity_config, dict):
            include_calculated_fields = entity_config.get('include_calculated_fields', [])
        
        # Check if we need to add any calculated fields
        for field_name in include_calculated_fields:
            if field_name not in data:
                # Add calculated field using existing patterns
                if field_name == 'amount_in_words':
                    # Get primary_amount_field safely
                    amount_field = None
                    if hasattr(entity_config, 'primary_amount_field'):
                        amount_field = entity_config.primary_amount_field
                    elif isinstance(entity_config, dict):
                        amount_field = entity_config.get('primary_amount_field')
                    
                    # Use existing number to words conversion
                    if amount_field and amount_field in data:
                        data[field_name] = self._convert_amount_to_words(data[amount_field])
                
                elif field_name.endswith('_name'):
                    # Relationship name fields
                    base_field = field_name[:-5]  # Remove '_name'
                    if base_field in data:
                        data[field_name] = self._get_related_name(base_field, data[base_field])
        
        return data
    


    # FIXED prepare_document_context method to handle entity_config as both dict and object

    def prepare_document_context(self, entity_config, doc_config, doc_data: Dict) -> Dict:
        """
        Prepare context for document template using enhanced context helpers
        """
        from flask_login import current_user
        from datetime import datetime
        from app.utils.context_helpers import get_document_organization_context
        
        # Extract actual data first
        actual_data = doc_data
        if isinstance(doc_data, dict):
            wrapper_keys = {'item', 'payment', 'entity_dict', 'data'}
            doc_keys = set(doc_data.keys())
            
            if len(doc_keys) == 1 and doc_keys.issubset(wrapper_keys):
                wrapper_key = list(doc_keys)[0]
                if isinstance(doc_data[wrapper_key], dict):
                    actual_data = doc_data[wrapper_key]
        
        # Get complete organization context (hospital + branch) from helper
        organization_data = get_document_organization_context()
        
        self.logger.info(f"Hospital: {organization_data['hospital']['name']} with logo: {organization_data['hospital'].get('logo_path')}")
        self.logger.info(f"Branch: {organization_data['branch']['name']}")
        
        # Simple assembled_data
        assembled_data = {
            'tabs': {},
            'layout_type': 'simple',
            'field_sections': [],
            'has_tabs': False
        }
        
        # Try to assemble view data if we have proper config
        try:
            from app.engine.data_assembler import EnhancedUniversalDataAssembler
            assembler = EnhancedUniversalDataAssembler()
            
            if hasattr(entity_config, '__dict__'):
                assembled_data = assembler.assemble_universal_view_data(
                    config=entity_config,
                    raw_item_data={'item': actual_data},
                    current_user=current_user,
                    branch_context=organization_data
                )
        except Exception as e:
            self.logger.debug(f"Using simple assembled_data: {e}")
        
        # Build context for template
        context = {
            'entity_config': entity_config,
            'doc_config': doc_config,
            'data': actual_data,
            'assembled_data': assembled_data,
            'organization': organization_data,
            'hospital': organization_data['hospital'],
            'branch': organization_data['branch'],
            'current_user': current_user,
            'current_datetime': datetime.now(),
            'hospital_name': organization_data['hospital']['name'],
            'branch_name': organization_data['branch']['name']
        }
        
        # Add entity_name
        entity_name = 'Document'
        if hasattr(entity_config, 'name'):
            entity_name = entity_config.name
        elif hasattr(entity_config, 'plural_name'):
            entity_name = entity_config.plural_name
        elif isinstance(entity_config, dict):
            entity_name = entity_config.get('name', entity_config.get('plural_name', 'Document'))
        
        context['entity_name'] = entity_name
        
        return context
    

    def _convert_amount_to_words(self, amount) -> str:
        """
        Convert amount to words (Indian numbering system)
        """
        try:
            from num2words import num2words
            
            # Convert to float if string
            if isinstance(amount, str):
                amount = float(amount.replace(',', ''))
            
            # Convert to words in Indian format
            words = num2words(amount, to='currency', lang='en_IN', currency='INR')
            return words.title()
        except ImportError:
            # Fallback if num2words not installed
            return f"Rupees {amount}"
        except Exception as e:
            self.logger.warning(f"Could not convert amount to words: {e}")
            return str(amount)
    
    def _get_related_name(self, field_name: str, field_value: Any) -> str:
        """
        Get related name for a field (e.g., supplier_name from supplier_id)
        This is a placeholder - actual implementation would fetch from database
        """
        # This would typically fetch from database
        # For now, return the value itself or empty string
        if field_value:
            return f"{field_name}_{field_value}"
        return ""
        
    
    def get_document_buttons(self, entity_config, entity_id: str) -> List[Dict]:
        """
        Generate document action buttons for view page
        Fixed to handle entity_config as dict or object
        """
        buttons = []
        
        # Get entity_type safely
        entity_type = None
        if hasattr(entity_config, 'entity_type'):
            entity_type = entity_config.entity_type
        elif isinstance(entity_config, dict):
            entity_type = entity_config.get('entity_type')
        
        if not entity_type:
            self.logger.warning("Could not determine entity_type for document buttons")
            return buttons
        
        # Get document_configs safely
        document_configs = {}
        if hasattr(entity_config, 'document_configs'):
            document_configs = entity_config.document_configs
        elif isinstance(entity_config, dict):
            document_configs = entity_config.get('document_configs', {})
        
        # Generate buttons for each document type
        for doc_type, doc_config in document_configs.items():
            # Check if enabled
            enabled = True
            if hasattr(doc_config, 'enabled'):
                enabled = doc_config.enabled
            elif isinstance(doc_config, dict):
                enabled = doc_config.get('enabled', True)
            
            if enabled:
                # Get title safely
                title = 'Document'
                if hasattr(doc_config, 'title'):
                    title = doc_config.title
                elif isinstance(doc_config, dict):
                    title = doc_config.get('title', 'Document')
                
                # Print button
                buttons.append({
                    'id': f'print_{doc_type}',
                    'label': f'Print {title}',
                    'icon': 'fas fa-print',
                    'url': f'/universal/{entity_type}/document/{entity_id}/{doc_type}?auto_print=true',
                    'class': 'btn-outline',
                    'target': '_blank'
                })
                
                # Get allowed_formats safely
                allowed_formats = []
                if hasattr(doc_config, 'allowed_formats'):
                    allowed_formats = doc_config.allowed_formats
                elif isinstance(doc_config, dict):
                    allowed_formats = doc_config.get('allowed_formats', [])
                
                # PDF button if PDF format is allowed
                pdf_allowed = False
                for fmt in allowed_formats:
                    fmt_str = fmt.lower() if isinstance(fmt, str) else str(fmt).lower()
                    if 'pdf' in fmt_str:
                        pdf_allowed = True
                        break
                
                if pdf_allowed:
                    buttons.append({
                        'id': f'pdf_{doc_type}',
                        'label': f'Download {title} PDF',
                        'icon': 'fas fa-file-pdf',
                        'url': f'/universal/{entity_type}/document/{entity_id}/{doc_type}?format=pdf',
                        'class': 'btn-outline',
                        'target': '_blank'
                    })

                # Excel button if Excel format is allowed
                excel_allowed = False
                for fmt in allowed_formats:
                    fmt_str = fmt.lower() if isinstance(fmt, str) else str(fmt).lower()
                    if 'excel' in fmt_str:
                        excel_allowed = True
                        break
                
                if excel_allowed:
                    buttons.append({
                        'id': f'excel_{doc_type}',
                        'label': f'Export {title} Excel',
                        'icon': 'fas fa-file-excel',
                        'url': f'/universal/{entity_type}/document/{entity_id}/{doc_type}?format=excel',
                        'class': 'btn-outline',
                        'target': '_blank'
                    })
                
                # Word button if Word format is allowed
                word_allowed = False
                for fmt in allowed_formats:
                    fmt_str = fmt.lower() if isinstance(fmt, str) else str(fmt).lower()
                    if 'word' in fmt_str:
                        word_allowed = True
                        break
                
                if word_allowed:
                    buttons.append({
                        'id': f'word_{doc_type}',
                        'label': f'Export {title} Word',
                        'icon': 'fas fa-file-word',
                        'url': f'/universal/{entity_type}/document/{entity_id}/{doc_type}?format=word',
                        'class': 'btn-outline',
                        'target': '_blank'
                    })
        
        return buttons
    
    def render_document_html(self, context: Dict) -> str:
        """
        Render document as HTML using universal template
        """
        return render_template('engine/universal_document.html', **context)
    
    def render_document_pdf(self, context: Dict, doc_config) -> Any:
        """
        Generate PDF using ReportLab - Fixed to use field_sections correctly
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
            from io import BytesIO
            from flask import request
            import os
            
            # Create PDF buffer
            buffer = BytesIO()
            
            # Create the PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=15*mm,
                leftMargin=15*mm,
                topMargin=15*mm,
                bottomMargin=15*mm,
                title=doc_config.get('title', 'Document')
            )
            
            elements = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                textColor=colors.HexColor('#000000'),
                spaceAfter=15,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            card_title_style = ParagraphStyle(
                'CardTitle',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=colors.HexColor('#ffffff'),
                spaceAfter=0,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000')
            )
            
            label_style = ParagraphStyle(
                'Label',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#666666'),
                fontName='Helvetica-Bold'
            )
            
            value_style = ParagraphStyle(
                'Value',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000')
            )
            
            # Extract data
            data = context.get('data', {})
            hospital = context.get('hospital', {})
            branch = context.get('branch', {})
            assembled_data = context.get('assembled_data', {})
            entity_config = context.get('entity_config')
            
            # Log what we have
            self.logger.info(f"PDF Gen - assembled_data keys: {list(assembled_data.keys())}")
            self.logger.info(f"PDF Gen - field_sections exists: {bool(assembled_data.get('field_sections'))}")
            self.logger.info(f"PDF Gen - field_sections count: {len(assembled_data.get('field_sections', []))}")
            
            # Define card header color
            CARD_HEADER_COLOR = colors.HexColor('#5b9bd5')  # Light blue
            
            # Helper functions
            def get_value_or_dash(value, is_amount=False):
                """Return formatted value or dash if empty - ALWAYS returns string"""
                # ✅ Handle None, empty, or 'None' string
                if value is None or value == '' or str(value).lower() == 'none':
                    return '-'
                
                # ✅ Handle amount formatting
                if is_amount:
                    try:
                        # Convert to float first, then format
                        float_val = float(value)
                        return f"₹{float_val:,.2f}"  # Removed space after ₹ for consistency
                    except (ValueError, TypeError):
                        return '-'
                
                # ✅ Handle dates
                if hasattr(value, 'strftime'):
                    try:
                        return value.strftime('%d/%m/%Y')
                    except:
                        return str(value)
                
                # ✅ Handle boolean values
                if isinstance(value, bool):
                    return 'Yes' if value else 'No'
                
                # ✅ Always ensure string return
                try:
                    return str(value)
                except:
                    return '-'
            
            def create_card(title, content_table):
                """Create a card with title bar and content"""
                title_para = Paragraph(title.upper(), card_title_style)
                title_table = Table([[title_para]], colWidths=[530])
                title_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), CARD_HEADER_COLOR),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ]))
                
                card_data = [[title_table], [content_table]]
                card = Table(card_data, colWidths=[530])
                card.setStyle(TableStyle([
                    ('BOX', (0, 0), (-1, -1), 1, colors.grey),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                ]))
                
                return card
            
            # 1. COMPANY HEADER WITH LOGO
            header_data = []
            header_row = []
            
            company_lines = []
            company_lines.append(f"<font size='14'><b>{hospital.get('name', 'Healthcare Facility')}</b></font>")
            
            if hospital.get('address'):
                company_lines.append(f"<font size='9'>{hospital['address']}</font>")
            
            location_parts = []
            if hospital.get('city'):
                location_parts.append(hospital['city'])
            if hospital.get('state'):
                location_parts.append(hospital['state'])
            if hospital.get('pincode'):
                location_parts.append(hospital['pincode'])
            if location_parts:
                company_lines.append(f"<font size='9'>{', '.join(location_parts)}</font>")
            
            contact_info = []
            if hospital.get('phone'):
                contact_info.append(f"Phone: {hospital['phone']}")
            if hospital.get('email'):
                contact_info.append(f"Email: {hospital['email']}")
            if contact_info:
                company_lines.append(f"<font size='9'>{' | '.join(contact_info)}</font>")
            
            if hospital.get('gst_number'):
                company_lines.append(f"<font size='9'>GST: {hospital['gst_number']}</font>")
            
            company_text = "<br/>".join(company_lines)
            header_row.append(Paragraph(company_text, normal_style))
            
            # Logo handling
            logo_added = False
            if hospital.get('hospital_id') and hospital.get('logo_path'):
                logo_attempts = [
                    os.path.join('app', 'static', 'uploads', 'hospital_logos', 
                                str(hospital['hospital_id']), str(hospital['logo_path'])),
                    os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos',
                                str(hospital['hospital_id']), str(hospital['logo_path'])),
                ]
                
                for logo_path in logo_attempts:
                    if os.path.exists(logo_path):
                        try:
                            img = Image(logo_path, width=100, height=100, kind='bound')
                            img.hAlign = 'RIGHT'
                            header_row.append(img)
                            logo_added = True
                            break
                        except:
                            pass
            
            if not logo_added:
                header_row.append('')
            
            header_data.append(header_row)
            header_table = Table(header_data, colWidths=[400, 130])
            header_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            elements.append(header_table)
            elements.append(Spacer(1, 10))
            
            # Line separator
            line_table = Table([['']], colWidths=[530])
            line_table.setStyle(TableStyle([
                ('LINEBELOW', (0, 0), (-1, 0), 1.5, colors.black),
            ]))
            elements.append(line_table)
            elements.append(Spacer(1, 15))
            
            # 2. DOCUMENT TITLE
            title_text = doc_config.get('header_text', doc_config.get('title', 'DOCUMENT'))
            elements.append(Paragraph(title_text.upper(), title_style))
            elements.append(Spacer(1, 15))
            
            # 3. BUILD CARDS FROM FIELD_SECTIONS (which are already assembled!)
            
            # Get doc_config filters
            visible_tabs = doc_config.get('visible_tabs', [])
            hidden_sections = doc_config.get('hidden_sections', [])
            
            # Get field_sections directly from assembled_data
            field_sections = assembled_data.get('field_sections', [])
            
            self.logger.info(f"PDF Gen - Processing {len(field_sections)} sections")
            
            # Process each section as a card
            for section_data in field_sections:
                if not isinstance(section_data, dict):
                    continue
                
                # Each section_data is a tab with sections inside
                tab_key = section_data.get('key', '')
                tab_label = section_data.get('label', '')
                sections = section_data.get('sections', [])
                
                # Check if this tab should be visible
                if visible_tabs and tab_key not in visible_tabs:
                    self.logger.debug(f"Skipping tab {tab_key} - not in visible_tabs")
                    continue
                
                # Process sections within this tab
                for section in sections:
                    if not isinstance(section, dict):
                        continue
                    
                    section_key = section.get('key', '')
                    section_title = section.get('title', 'Information')
                    fields = section.get('fields', [])
                    
                    # Skip hidden sections
                    if hidden_sections and section_key in hidden_sections:
                        self.logger.debug(f"Skipping section {section_key} - in hidden_sections")
                        continue
                    
                    # Skip empty sections
                    if not fields:
                        continue
                    
                    # Process fields including custom renderers
                    # Process fields including custom renderers
                    has_custom_table = False
                    
                    # Check each field in the section
                    for field in fields:
                        # Check if this field has a custom renderer with table data
                        if field.get('is_custom_renderer') and field.get('custom_renderer'):
                            css_classes = field.get('custom_renderer', {}).get('css_classes', '')
                            
                            # Check if it's a table renderer (generic check)
                            if 'table-responsive' in css_classes:
                                field_label = field.get('label', 'Items')
                                
                                # Try to get data from field first
                                field_data = field.get('data')
                                
                                # ✅ FIX: For PDF, always fetch fresh data using service
                                if not field_data or not field_data.get('items'):  # ✅ Added check for empty items
                                    context_function = field.get('custom_renderer', {}).get('context_function')
                                    if context_function:
                                        # Get the service and call the function
                                        from app.engine.universal_services import get_universal_service
                                        entity_type = context.get('entity_type', '')
                                        service = get_universal_service(entity_type)
                                        
                                        if service and hasattr(service, context_function):
                                            try:
                                                method = getattr(service, context_function)
                                                field_data = method(
                                                    context.get('item_id'),
                                                    hospital_id=context.get('current_hospital_id'),
                                                    branch_id=context.get('current_branch_id')
                                                )
                                                self.logger.info(f"✅ PDF: Fetched {len(field_data.get('items', []))} items via {context_function}")
                                            except Exception as e:
                                                self.logger.error(f"PDF: Error calling {context_function}: {str(e)}")
                                                field_data = None
                                    if context_function:
                                        # Get the service and call the function
                                        from app.engine.universal_services import get_universal_service
                                        entity_type = context.get('entity_type', '')
                                        service = get_universal_service(entity_type)
                                        
                                        if service and hasattr(service, context_function):
                                            try:
                                                method = getattr(service, context_function)
                                                field_data = method(
                                                    context.get('item_id'),
                                                    hospital_id=context.get('current_hospital_id'),
                                                    branch_id=context.get('current_branch_id')
                                                )
                                            except Exception as e:
                                                self.logger.error(f"PDF: Error calling {context_function}: {str(e)}")
                                                field_data = None
                                
                                # Process data if available (GENERIC - no entity-specific logic)
                                if field_data and isinstance(field_data, dict) and field_data.get('items'):
                                    items = field_data['items']
                                    
                                    if items:
                                        # Dynamically determine columns from first item
                                        first_item = items[0]
                                        
                                        # Build columns dynamically from data keys
                                        total_width = 460  # Available width in PDF
                                        columns = []
                                        
                                        # Create columns from first item's keys
                                        for key in first_item.keys():
                                            # Generate label from key
                                            label = key.replace('_', ' ').title()
                                            
                                            # Estimate column width based on content type
                                            if 'name' in key or 'description' in key:
                                                width = 120  # Wider for text
                                            elif 'amount' in key or 'total' in key:
                                                width = 70   # Medium for amounts
                                            elif 'percent' in key or 'qty' in key or 'quantity' in key:
                                                width = 50   # Smaller for numbers
                                            else:
                                                width = 60   # Default
                                            
                                            columns.append((key, label, width))
                                        
                                        # Adjust widths to fit total width
                                        total_estimated = sum(w for _, _, w in columns)
                                        if total_estimated > 0:
                                            scale_factor = total_width / total_estimated
                                            columns = [(k, l, w * scale_factor) for k, l, w in columns]
                                        
                                        # Build header row
                                        header_row = []
                                        for col_key, col_label, _ in columns:
                                            header_row.append(Paragraph(f'<b>{col_label}</b>', label_style))
                                        
                                        table_rows = [header_row]
                                        
                                        # Build data rows
                                        for item in items:
                                            row = []
                                            for col_key, _, _ in columns:
                                                value = item.get(col_key, '')
                                                
                                                # ✅ FIX: Ensure value is never None or empty before processing
                                                if value is None or value == '':
                                                    value = '-'
                                                
                                                # Generic formatting based on key patterns
                                                if value and value != '-':
                                                    if any(pattern in col_key for pattern in ['amount', 'price', 'total', 'cost']):
                                                        value = get_value_or_dash(value, is_amount=True)
                                                    elif 'percent' in col_key:
                                                        try:
                                                            value = f"{float(value):.1f}%"
                                                        except:
                                                            value = str(value)
                                                    else:
                                                        value = str(value)
                                                
                                                # ✅ CRITICAL FIX: Always ensure value is a string before Paragraph()
                                                if not isinstance(value, str):
                                                    value = str(value) if value is not None else '-'
                                                
                                                row.append(Paragraph(value, value_style))
                                            table_rows.append(row)
                                        
                                        # Add summary row if exists (generic handling)
                                        if field_data.get('summary'):
                                            summary = field_data['summary']
                                            summary_row = []
                                            
                                            # First column shows "TOTAL"
                                            summary_row.append(Paragraph('<b>TOTAL</b>', label_style))
                                            
                                            # Other columns show summary values or empty
                                            for i, (col_key, _, _) in enumerate(columns[1:], 1):
                                                # Check if summary has this column
                                                if col_key in summary:
                                                    value = summary[col_key]
                                                    if any(pattern in col_key for pattern in ['amount', 'price', 'total', 'gst']):
                                                        value = get_value_or_dash(value, is_amount=True)
                                                    else:
                                                        value = str(value)
                                                    summary_row.append(Paragraph(f'<b>{value}</b>', label_style))
                                                else:
                                                    summary_row.append('')
                                            
                                            table_rows.append(summary_row)
                                        
                                        # Create table with calculated column widths
                                        # ✅ FORCE TABLE TO USE FULL WIDTH: Calculate proportional widths
                                        total_proportional_width = sum(w for _, _, w in columns)
                                        available_width = 530  # Maximum content width in A4 (210mm - 30mm margins = 180mm ≈ 530pt)

                                        # Calculate actual column widths as proportions of available width
                                        col_widths = []
                                        for _, _, prop_width in columns:
                                            actual_width = (prop_width / total_proportional_width) * available_width
                                            col_widths.append(actual_width)

                                        # ✅ Verify total width equals available width
                                        total_calc_width = sum(col_widths)
                                        if abs(total_calc_width - available_width) > 1:  # Allow 1pt tolerance
                                            # Adjust last column to exactly match available width
                                            col_widths[-1] = col_widths[-1] + (available_width - total_calc_width)
                                        items_table = Table(table_rows, colWidths=col_widths)
                                        
                                        # Apply generic styling
                                        items_table.setStyle(TableStyle([
                                            # Header
                                            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                            ('FONTSIZE', (0, 0), (-1, 0), 9),
                                            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                                            
                                            # Body
                                            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                            ('FONTSIZE', (0, 1), (-1, -1), 8),
                                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                            ('LEFTPADDING', (0, 0), (-1, -1), 5),
                                            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                                            ('TOPPADDING', (0, 0), (-1, -1), 4),
                                            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                                        ]))
                                        
                                        # Dynamic alignment based on content
                                        for i, (col_key, _, _) in enumerate(columns):
                                            # Left align text columns, right align numeric
                                            if any(pattern in col_key for pattern in ['name', 'description', 'item', 'batch']):
                                                items_table.setStyle(TableStyle([
                                                    ('ALIGN', (i, 1), (i, -1), 'LEFT'),
                                                ]))
                                            else:
                                                items_table.setStyle(TableStyle([
                                                    ('ALIGN', (i, 1), (i, -1), 'RIGHT'),
                                                ]))
                                        
                                        # Summary row styling if exists
                                        if field_data.get('summary'):
                                            items_table.setStyle(TableStyle([
                                                ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
                                                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                                                ('LINEABOVE', (0, -1), (-1, -1), 1, colors.black),
                                            ]))
                                        
                                        # Create card using existing helper
                                        table_card = create_card(field_label, items_table)
                                        elements.append(table_card)
                                        elements.append(Spacer(1, 15))
                                        
                                has_custom_table = True
                                break  # Only process one custom renderer per section
                    
                    # Skip regular field processing if we handled a custom table
                    if has_custom_table:
                        continue
                    
                    # Build rows for this section
                    section_rows = []
                    field_count = 0
                    current_row = []
                    
                    for field in fields:
                        # Get field details
                        field_label = field.get('label', '')
                        field_value = field.get('value')
                        field_key = field.get('key', '')
                        
                        # Format value based on type or key
                        if field_value is None:
                            field_value = '-'
                        elif 'amount' in field_key.lower() or field.get('type') == 'currency':
                            field_value = get_value_or_dash(field_value, is_amount=True)
                        elif 'date' in field_key.lower() or field.get('type') == 'date':
                            field_value = get_value_or_dash(field_value)
                        else:
                            field_value = get_value_or_dash(field_value)
                        
                        # Add to current row
                        current_row.extend([
                            Paragraph(f"{field_label}:", label_style),
                            Paragraph(field_value, value_style)
                        ])
                        field_count += 1
                        
                        # Start new row after 2 fields
                        if field_count % 2 == 0:
                            section_rows.append(current_row)
                            current_row = []
                    
                    # Add remaining fields
                    if current_row:
                        while len(current_row) < 4:
                            current_row.append('')
                        section_rows.append(current_row)
                    
                    # Create table for this section
                    if section_rows:
                        section_table = Table(section_rows, colWidths=[85, 165, 85, 165])
                        section_table.setStyle(TableStyle([
                            ('LEFTPADDING', (0, 0), (-1, -1), 10),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ]))
                        
                        # Create and add card
                        section_card = create_card(section_title, section_table)
                        elements.append(section_card)
                        elements.append(Spacer(1, 15))
            
            # 4. TERMS AND CONDITIONS
            if doc_config.get('show_terms') and doc_config.get('terms_content'):
                terms_content = []
                terms = doc_config.get('terms_content', [])
                
                if isinstance(terms, list):
                    for i, term in enumerate(terms, 1):
                        terms_content.append(Paragraph(f"{i}. {term}", value_style))
                        terms_content.append(Spacer(1, 5))
                else:
                    terms_content.append(Paragraph(terms, value_style))
                
                terms_table = Table([[terms_content]], colWidths=[500])
                terms_table.setStyle(TableStyle([
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                terms_card = create_card('TERMS AND CONDITIONS', terms_table)
                elements.append(terms_card)
                elements.append(Spacer(1, 20))
            
            # 5. SIGNATURE SECTION
            if doc_config.get('signature_fields'):
                elements.append(Spacer(1, 30))
                
                sig_fields = []
                for sig_field in doc_config.get('signature_fields', []):
                    sig_fields.append([
                        Spacer(1, 30),
                        Table([['']], colWidths=[150], rowHeights=[1],
                            style=TableStyle([('LINEBELOW', (0, 0), (-1, 0), 1, colors.black)])),
                        Paragraph(sig_field.get('label', 'Signature'), label_style)
                    ])
                
                sig_cols = len(sig_fields)
                if sig_cols > 0:
                    col_width = 530 / sig_cols
                    sig_table = Table([sig_fields], colWidths=[col_width] * sig_cols)
                    sig_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    
                    elements.append(sig_table)
                    elements.append(Spacer(1, 20))
            
            # 6. FOOTER
            if doc_config.get('show_footer'):
                footer_lines = []
                footer_lines.append(doc_config.get('footer_text', 'This is a computer generated document'))
                
                if doc_config.get('show_print_info'):
                    from datetime import datetime
                    footer_lines.append(f"Generated on: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
                
                footer_text = "<br/>".join([f"<i>{line}</i>" for line in footer_lines])
                footer_para = Paragraph(footer_text, label_style)
                footer_para.alignment = TA_CENTER
                
                elements.append(Spacer(1, 10))
                elements.append(footer_para)
            
            # Build PDF
            doc.build(elements)
            
            buffer.seek(0)
            response = make_response(buffer.read())
            response.headers['Content-Type'] = 'application/pdf'
            
            # Generate filename
            title_field = 'reference_no'
            if entity_config and hasattr(entity_config, 'title_field'):
                title_field = entity_config.title_field
            
            doc_ref = data.get(title_field, 'document')
            filename = f"{doc_config.get('title', 'Document')}_{doc_ref}.pdf"
            
            disposition = 'attachment' if request.args.get('download') == 'true' else 'inline'
            response.headers['Content-Disposition'] = f'{disposition}; filename="{filename}"'
            
            return response
            
        except Exception as e:
            self.logger.error(f"ReportLab PDF generation failed: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return self.render_document_html(context)
        
    # ========== NEW METHOD 1: Excel Renderer ==========
    def render_document_excel(self, context: Dict, doc_config: Dict) -> Any:
        """
        NEW METHOD - Render document as Excel
        Uses same context structure as existing render_document_pdf
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
            from io import BytesIO
            
            # Extract data using SAME structure as render_document_pdf
            data = context.get('data', {})
            entity_config = context.get('entity_config')
            entity_type = context.get('entity_type')
            tabs = context.get('tabs', {})
            sections = context.get('sections', {})
            item = context.get('item', data)  # Same fallback as PDF
            item_id = context.get('item_id')
            hospital = context.get('hospital', {})
            branch = context.get('branch', {})
            
            # Create workbook
            wb = Workbook()
            ws = wb.active
            ws.title = doc_config.get('title', 'Document')[:31]  # Excel sheet name limit
            
            current_row = 1
            
            # Define styles (similar to PDF styles)
            header_font = Font(size=16, bold=True)
            title_font = Font(size=14, bold=True)
            section_font = Font(size=12, bold=True)
            label_font = Font(bold=True)
            header_fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
            
            # HEADER SECTION (matching PDF structure)
            if doc_config.get('show_company_info'):
                # Company name
                ws.merge_cells(f'A{current_row}:F{current_row}')
                cell = ws[f'A{current_row}']
                cell.value = hospital.get('name', 'Skinspire Clinic')
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center')
                current_row += 1
                
                # Branch name if exists
                if branch and branch.get('name'):
                    ws.merge_cells(f'A{current_row}:F{current_row}')
                    cell = ws[f'A{current_row}']
                    cell.value = branch['name']
                    cell.alignment = Alignment(horizontal='center')
                    current_row += 1
                
                current_row += 1  # Empty row
            
            # LOGO HANDLING for Excel (Add after the header section, before Document title)
            if doc_config.get('show_logo') and hospital.get('hospital_id'):
                try:
                    from openpyxl.drawing.image import Image as ExcelImage
                    
                    # Build logo search patterns with ABSOLUTE paths
                    logo_patterns = []
                    hospital_id_str = str(hospital['hospital_id'])

                    if hospital.get('logo_path'):
                        logo_path_str = str(hospital['logo_path'])
                        
                        # If it's an icon, try to find full-size version first
                        if 'icon' in logo_path_str.lower():
                            # Extract base filename (remove 'icon_' prefix)
                            base_name = logo_path_str.replace('icon_', '')
                            
                            # Try full-size version first
                            logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                            hospital_id_str, base_name))
                            # Also try with 'logo_' prefix
                            logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                            hospital_id_str, 'logo_' + base_name))
                        else:
                            # Non-icon version, use as-is
                            logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                            hospital_id_str, logo_path_str))

                    # Add standard patterns with absolute paths
                    for filename in ['logo_large.png', 'logo_large.jpg', 'logo.png', 'logo.jpg']:
                        logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                        hospital_id_str, filename))

                    # Finally add the icon itself as fallback (if it exists)
                    if hospital.get('logo_path') and 'icon' in str(hospital.get('logo_path')).lower():
                        logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                        hospital_id_str, str(hospital['logo_path'])))
                    
                    logo_added = False
                    for logo_path in logo_patterns:
                        if logo_path and os.path.exists(logo_path):
                            self.logger.info(f"Logo file exists at: {logo_path}, size: {os.path.getsize(logo_path)} bytes")
                            try:
                                self.logger.debug(f"Trying to load logo from: {logo_path}")
                                img = ExcelImage(logo_path)
                                # Resize logo to fit
                                img.height = 80
                                img.width = 80
                                # Place logo in cell G1 to avoid merge conflicts
                                ws.add_image(img, 'G1')
                                logo_added = True
                                self.logger.info(f"Successfully added logo to Excel from: {logo_path}")
                                break
                            except Exception as e:
                                self.logger.debug(f"Failed to load logo from {logo_path}: {e}")
                                continue
                    
                    if not logo_added:
                        self.logger.debug("No logo found for Excel document")
                        
                except ImportError:
                    self.logger.warning("openpyxl Image support not available - install Pillow: pip install Pillow")
                except Exception as e:
                    self.logger.error(f"Error adding logo to Excel: {e}")

            current_row += 1  # Add extra spacing after logo area    

            # Document title
            ws.merge_cells(f'A{current_row}:F{current_row}')
            cell = ws[f'A{current_row}']
            cell.value = doc_config.get('title', 'Document')
            cell.font = title_font
            cell.alignment = Alignment(horizontal='center')
            current_row += 2
            
            # Metadata
            ws[f'A{current_row}'] = 'Generated Date:'
            ws[f'A{current_row}'].font = label_font
            ws[f'B{current_row}'] = datetime.now().strftime('%d/%m/%Y %H:%M')
            current_row += 1
            
            ws[f'A{current_row}'] = 'Generated By:'
            ws[f'A{current_row}'].font = label_font
            ws[f'B{current_row}'] = current_user.full_name if current_user else 'System'
            current_row += 2
            
            # BODY SECTIONS - Process field_sections (same as PDF)
            assembled_data = context.get('assembled_data', {})
            field_sections = assembled_data.get('field_sections', [])
            visible_tabs = doc_config.get('visible_tabs', [])
            hidden_sections = doc_config.get('hidden_sections', [])

            self.logger.info(f"Excel Gen - Processing {len(field_sections)} sections")

            for section_data in field_sections:
                if not isinstance(section_data, dict):
                    continue
                
                # Each section_data is a tab with sections inside
                tab_key = section_data.get('key', '')
                tab_label = section_data.get('label', '')
                sections = section_data.get('sections', [])
                
                # Check if this tab should be visible
                if visible_tabs and tab_key not in visible_tabs:
                    self.logger.debug(f"Skipping tab {tab_key} - not in visible_tabs")
                    continue
                
                # Tab title
                ws[f'A{current_row}'] = tab_label
                ws[f'A{current_row}'].font = section_font
                ws[f'A{current_row}'].fill = header_fill
                ws.merge_cells(f'A{current_row}:F{current_row}')
                current_row += 1
                
                # Process sections within tab
                for section in sections:
                    if not isinstance(section, dict):
                        continue
                    
                    section_key = section.get('key', '')
                    section_title = section.get('title', 'Information')
                    fields = section.get('fields', [])
                    
                    # Skip hidden sections
                    if hidden_sections and section_key in hidden_sections:
                        self.logger.debug(f"Skipping section {section_key} - in hidden_sections")
                        continue
                    
                    # Skip empty sections
                    if not fields:
                        continue
                    
                    # Section title
                    if section_title:
                        ws[f'A{current_row}'] = section_title
                        ws[f'A{current_row}'].font = Font(bold=True, size=11)
                        current_row += 1
                    
                    # Section fields
                    for field in fields:
                        if not isinstance(field, dict):
                            continue
                        
                        field_label = field.get('label', '')
                        field_value = field.get('value', '')
                        field_type = field.get('type', 'text')
                        
                        # Skip empty fields
                        if not field_label:
                            continue
                        
                        # Label
                        ws[f'A{current_row}'] = str(field_label)
                        ws[f'A{current_row}'].font = label_font
                        
                        # Value formatting
                        if field_type in ['currency', 'amount'] and field_value:
                            try:
                                # Try to convert to number for Excel formatting
                                if isinstance(field_value, str) and '₹' in field_value:
                                    field_value = field_value.replace('₹', '').replace(',', '').strip()
                                ws[f'B{current_row}'] = float(field_value) if field_value else 0
                                ws[f'B{current_row}'].number_format = '₹#,##0.00'
                            except:
                                ws[f'B{current_row}'] = str(field_value)
                        elif field_type == 'date' and field_value:
                            ws[f'B{current_row}'] = str(field_value)
                            ws[f'B{current_row}'].number_format = 'DD/MM/YYYY'
                        else:
                            ws[f'B{current_row}'] = str(field_value) if field_value else ''
                        
                        current_row += 1
                    
                    current_row += 1  # Empty row between sections
            
            # LINE ITEMS (if configured, same as PDF)
            if doc_config.get('show_line_items'):
                current_row += 1
                
                # Check for line items in data (same locations as PDF checks)
                line_items = []
                if 'linked_invoices' in item:
                    line_items = item['linked_invoices']
                elif 'line_items' in item:
                    line_items = item['line_items']
                elif 'items' in item:
                    line_items = item['items']
                
                if line_items and len(line_items) > 0:
                    # Title
                    ws[f'A{current_row}'] = 'Details'
                    ws[f'A{current_row}'].font = section_font
                    ws[f'A{current_row}'].fill = header_fill
                    ws.merge_cells(f'A{current_row}:F{current_row}')
                    current_row += 1
                    
                    # Get columns from config
                    columns = doc_config.get('line_item_columns', [])
                    if columns:
                        # Headers
                        for col_idx, col_name in enumerate(columns, 1):
                            cell = ws.cell(row=current_row, column=col_idx)
                            cell.value = col_name.replace('_', ' ').title()
                            cell.font = label_font
                            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                            cell.border = Border(
                                bottom=Side(style='thin'),
                                top=Side(style='thin'),
                                left=Side(style='thin'),
                                right=Side(style='thin')
                            )
                        current_row += 1
                        
                        # Data rows
                        for item_row in line_items:
                            for col_idx, col_name in enumerate(columns, 1):
                                cell = ws.cell(row=current_row, column=col_idx)
                                cell.value = str(item_row.get(col_name, ''))
                                cell.border = Border(
                                    bottom=Side(style='thin'),
                                    left=Side(style='thin'),
                                    right=Side(style='thin')
                                )
                            current_row += 1
            
            # FOOTER (if configured)
            if doc_config.get('show_footer'):
                current_row += 2
                
                # Footer text
                if doc_config.get('footer_text'):
                    ws.merge_cells(f'A{current_row}:F{current_row}')
                    ws[f'A{current_row}'] = doc_config['footer_text']
                    ws[f'A{current_row}'].alignment = Alignment(horizontal='center')
                    current_row += 1
                
                # Print info
                if doc_config.get('show_print_info'):
                    ws.merge_cells(f'A{current_row}:F{current_row}')
                    ws[f'A{current_row}'] = f"Generated on {datetime.now().strftime('%d/%m/%Y %H:%M')}"
                    ws[f'A{current_row}'].alignment = Alignment(horizontal='center')
                    ws[f'A{current_row}'].font = Font(size=9, italic=True)
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                
                for cell in column:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save to buffer
            excel_buffer = BytesIO()
            wb.save(excel_buffer)
            excel_buffer.seek(0)
            
            # Create response (same pattern as PDF)
            response = make_response(excel_buffer.read())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            
            # Generate filename (same pattern as PDF)
            title_field = 'reference_no'
            if entity_config and hasattr(entity_config, 'title_field'):
                title_field = entity_config.title_field
            
            doc_ref = item.get(title_field, 'document')
            filename = f"{doc_config.get('title', 'Document')}_{doc_ref}_{datetime.now().strftime('%Y%m%d')}.xlsx"
            
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as e:
            self.logger.error(f"Excel generation failed: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            # Fallback to HTML (same as PDF does)
            return self.render_document_html(context)
    
    # ========== NEW METHOD 2: Word Renderer ==========
    def render_document_word(self, context: Dict, doc_config: Dict) -> Any:
        """
        NEW METHOD - Render document as Word
        Uses same context structure as existing render_document_pdf
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from io import BytesIO
            
            # Extract data using SAME structure as render_document_pdf
            data = context.get('data', {})
            entity_config = context.get('entity_config')
            entity_type = context.get('entity_type')
            tabs = context.get('tabs', {})
            sections = context.get('sections', {})
            item = context.get('item', data)
            item_id = context.get('item_id')
            hospital = context.get('hospital', {})
            branch = context.get('branch', {})
            
            # Create document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = doc_config.get('title', 'Document')
            doc.core_properties.author = current_user.full_name if current_user else 'System'
            
            # HEADER SECTION with LOGO (matching PDF structure)
            if doc_config.get('show_company_info'):
                # Create a table for header with logo
                header_table = doc.add_table(rows=1, cols=2)
                header_table.autofit = False
                header_table.allow_autofit = False
                
                # Left cell - Company details
                left_cell = header_table.rows[0].cells[0]
                left_cell.width = Inches(4.5)
                
                # Company name
                company_para = left_cell.add_paragraph()
                company_run = company_para.add_run(hospital.get('name', 'Skinspire Clinic'))
                company_run.font.size = Pt(16)
                company_run.font.bold = True
                
                # Branch name if exists
                if branch and branch.get('name'):
                    branch_para = left_cell.add_paragraph()
                    branch_run = branch_para.add_run(branch['name'])
                    branch_run.font.size = Pt(12)
                
                # Address and contact info (optional)
                if hospital.get('address'):
                    addr_para = left_cell.add_paragraph()
                    addr_run = addr_para.add_run(hospital['address'])
                    addr_run.font.size = Pt(10)
                
                # Right cell - Logo
                right_cell = header_table.rows[0].cells[1]
                right_cell.width = Inches(2)
                
                # LOGO HANDLING for Word
                if doc_config.get('show_logo') and hospital.get('hospital_id'):
                    # Build logo search patterns with ABSOLUTE paths
                    logo_patterns = []
                    hospital_id_str = str(hospital['hospital_id'])

                    if hospital.get('logo_path'):
                        logo_path_str = str(hospital['logo_path'])
                        
                        # If it's an icon, try to find full-size version first
                        if 'icon' in logo_path_str.lower():
                            # Extract base filename (remove 'icon_' prefix)
                            base_name = logo_path_str.replace('icon_', '')
                            
                            # Try full-size version first
                            logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                            hospital_id_str, base_name))
                            # Also try with 'logo_' prefix
                            logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                            hospital_id_str, 'logo_' + base_name))
                        else:
                            # Non-icon version, use as-is
                            logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                            hospital_id_str, logo_path_str))

                    # Add standard patterns with absolute paths
                    for filename in ['logo_large.png', 'logo_large.jpg', 'logo.png', 'logo.jpg']:
                        logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                        hospital_id_str, filename))

                    # Finally add the icon itself as fallback (if it exists)
                    if hospital.get('logo_path') and 'icon' in str(hospital.get('logo_path')).lower():
                        logo_patterns.append(os.path.join(os.getcwd(), 'app', 'static', 'uploads', 'hospital_logos', 
                                                        hospital_id_str, str(hospital['logo_path'])))
                    
                    logo_added = False
                    for logo_path in logo_patterns:
                        if logo_path and os.path.exists(logo_path):
                            self.logger.info(f"Logo file exists at: {logo_path}, size: {os.path.getsize(logo_path)} bytes")
                            try:
                                self.logger.debug(f"Trying to load logo from: {logo_path}")
                                logo_para = right_cell.paragraphs[0]
                                logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                logo_run = logo_para.add_run()
                                logo_run.add_picture(logo_path, width=Inches(1.2))
                                logo_added = True
                                self.logger.info(f"Successfully added logo to Word from: {logo_path}")
                                break
                            except Exception as e:
                                self.logger.debug(f"Failed to load logo from {logo_path}: {e}")
                                continue
                    
                    if not logo_added:
                        self.logger.debug("No logo found for Word document")
                
                # Remove table borders
                for row in header_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = Pt(0)
                
                doc.add_paragraph()  # Empty line after header
            
            # Document title
            title = doc.add_heading(doc_config.get('title', 'Document'), 1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            metadata = doc.add_paragraph()
            metadata.add_run('Generated: ').bold = True
            metadata.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
            metadata.add_run('  |  ')
            metadata.add_run('By: ').bold = True
            doc.core_properties.author = current_user.full_name if current_user else 'System'
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Empty line
            
            # BODY SECTIONS - Process field_sections (same as PDF)
            assembled_data = context.get('assembled_data', {})
            field_sections = assembled_data.get('field_sections', [])
            visible_tabs = doc_config.get('visible_tabs', [])
            hidden_sections = doc_config.get('hidden_sections', [])

            self.logger.info(f"Word Gen - Processing {len(field_sections)} sections")

            for section_data in field_sections:
                if not isinstance(section_data, dict):
                    continue
                
                # Each section_data is a tab with sections inside
                tab_key = section_data.get('key', '')
                tab_label = section_data.get('label', '')
                sections = section_data.get('sections', [])
                
                # Check if this tab should be visible
                if visible_tabs and tab_key not in visible_tabs:
                    self.logger.debug(f"Skipping tab {tab_key} - not in visible_tabs")
                    continue
                
                # Tab as heading
                if tab_label:
                    doc.add_heading(tab_label, 2)
                
                # Process sections within tab
                for section in sections:
                    if not isinstance(section, dict):
                        continue
                    
                    section_key = section.get('key', '')
                    section_title = section.get('title', 'Information')
                    fields = section.get('fields', [])
                    
                    # Skip hidden sections
                    if hidden_sections and section_key in hidden_sections:
                        self.logger.debug(f"Skipping section {section_key} - in hidden_sections")
                        continue
                    
                    # Skip empty sections
                    if not fields:
                        continue
                    
                    # Section title
                    if section_title:
                        doc.add_heading(section_title, 3)
                    
                    # Create table for fields
                    if fields:
                        table = doc.add_table(rows=0, cols=2)
                        table.style = 'Light Grid Accent 1'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        for field in fields:
                            if not isinstance(field, dict):
                                continue
                                
                            field_label = field.get('label', '')
                            field_value = field.get('value', '')
                            
                            # Skip empty fields
                            if not field_label:
                                continue
                            
                            # Add row to table
                            row = table.add_row()
                            row.cells[0].text = str(field_label)
                            
                            # Format value
                            if field_value is None:
                                value_text = '-'
                            elif isinstance(field_value, bool):
                                value_text = 'Yes' if field_value else 'No'
                            elif isinstance(field_value, (list, dict)):
                                value_text = str(field_value)
                            else:
                                value_text = str(field_value)
                            
                            row.cells[1].text = value_text
                        
                        # Add spacing after table
                        doc.add_paragraph()
            
            # LINE ITEMS (if configured, same as PDF)
            if doc_config.get('show_line_items'):
                # Check for line items in data
                line_items = []
                if 'linked_invoices' in item:
                    line_items = item['linked_invoices']
                elif 'line_items' in item:
                    line_items = item['line_items']
                elif 'items' in item:
                    line_items = item['items']
                
                if line_items and len(line_items) > 0:
                    doc.add_heading('Details', 2)
                    
                    columns = doc_config.get('line_item_columns', [])
                    if columns:
                        table = doc.add_table(rows=1, cols=len(columns))
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        # Headers
                        for idx, col_name in enumerate(columns):
                            cell = table.rows[0].cells[idx]
                            cell.text = col_name.replace('_', ' ').title()
                            # Make header bold
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        
                        # Data rows
                        for item_row in line_items:
                            row = table.add_row()
                            for idx, col_name in enumerate(columns):
                                row.cells[idx].text = str(item_row.get(col_name, ''))
                        
                        doc.add_paragraph()
            
            # FOOTER (if configured)
            if doc_config.get('show_footer'):
                if doc_config.get('footer_text'):
                    doc.add_paragraph()  # Space before footer
                    footer = doc.add_paragraph(doc_config['footer_text'])
                    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Signatures if configured
                signature_fields = doc_config.get('signature_fields', [])
                if signature_fields:
                    doc.add_paragraph()
                    doc.add_paragraph()
                    
                    # Create signature table
                    sig_table = doc.add_table(rows=2, cols=len(signature_fields))
                    
                    for idx, sig in enumerate(signature_fields):
                        # Signature line
                        sig_table.rows[0].cells[idx].text = '_' * 30
                        # Signature label
                        sig_table.rows[1].cells[idx].text = sig.get('label', '')
                        
                        # Center align
                        for row in sig_table.rows:
                            row.cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to buffer
            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)
            
            # Create response (same pattern as PDF)
            response = make_response(word_buffer.read())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # Generate filename (same pattern as PDF)
            title_field = 'reference_no'
            if entity_config and hasattr(entity_config, 'title_field'):
                title_field = entity_config.title_field
            
            doc_ref = item.get(title_field, 'document')
            filename = f"{doc_config.get('title', 'Document')}_{doc_ref}_{datetime.now().strftime('%Y%m%d')}.docx"
            
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except ImportError:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            # Fallback to HTML
            return self.render_document_html(context)
        except Exception as e:
            self.logger.error(f"Word generation failed: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            # Fallback to HTML (same as PDF does)
            return self.render_document_html(context)
    

# Create singleton instance
document_service = UniversalDocumentService()

# Export functions for use in views
def get_document_service():
    """Get the document service instance"""
    return document_service

def prepare_document_data(entity_type: str, entity_id: str, data: Dict):
    """Prepare and store document data"""
    return document_service.store_document_data_in_session(entity_type, entity_id, data)

def get_document_buttons(entity_config, entity_id: str):
    """Get document action buttons"""
    return document_service.get_document_buttons(entity_config, entity_id)